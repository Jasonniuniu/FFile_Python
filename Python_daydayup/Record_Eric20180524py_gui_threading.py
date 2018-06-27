# -*- coding: utf-8 -*-

"""
Module implementing MainWindow.
"""
from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.QtCore import pyqtSlot
from PyQt5.QtWidgets import QMainWindow
from PyQt5.QtWidgets import QMessageBox  
from PyQt5.QtWidgets import QLineEdit,QInputDialog
from PyQt5.QtWidgets import QFileDialog  

from Ui_Eric20180524py_gui import Ui_MainWindow
from PyQt5.QtWidgets import QDialog
from Ui_Eric20180525Info import Ui_dialog

import webbrowser

import win32com.client    #调用系统的声音读取文本

import docx

from PyQt5.QtCore import *
from PyQt5.Qt import QSplashScreen 
from PyQt5.QtGui import QPixmap 
import time

#ASR
import wave
from pyaudio import PyAudio,paInt16
import json
import base64
import os
import requests
import time
#多线程
import threading
import queue as Queue
from time import sleep, ctime


def niu_read_docx(filename):
     doc=docx.Document(filename)
     fulltext=[]
     for para in doc.paragraphs:
         fulltext.append(para.text)
     return '\n'.join(fulltext)
         



class dialog(QDialog, Ui_dialog):
    """
    Class documentation goes here.
    """
    def __init__(self, parent=None):
        """
        Constructor
        
        @param parent reference to the parent widget
        @type QWidget
        """
        super(dialog, self).__init__(parent)
        self.setupUi(self)
        time.sleep(1) #使开机画面停留2秒 

    @pyqtSlot()
    def on_pushButton_clicked(self):
        """
        Slot documentation goes here.
        """
        my_str=self.lineEdit.text()      #获取单行文本框的内容
        my_str_2=self.lineEdit_2.text()
        my_str_3=self.lineEdit_3.text()
        my_str_4=self.lineEdit_4.text()
        
        print(my_str, my_str_2, my_str_3, my_str_4)
        Button1_1=QMessageBox.information(self, u'提示信息框', u'输入的s所有信息已经存储到数据库中！', ) #提示对话信息框
        print('OK')
        self.close() #关闭当前界面

    
    @pyqtSlot()
    def on_pushButton_2_clicked(self):
        """
        Slot documentation goes here.
        """
        print('Cancel')



class MyThread(threading.Thread):   #自定义多线程的类
    def __init__(seLf, func, args, name=''):  
        threading.Thread.__init__(seLf)  #对类进行初始化
        seLf .name=name                  #传递三个参数
        seLf.func=func
        seLf .args=args 
        
    def getresult(seLf):  #获得结果
        return self.res
        
    def run(self):  #重写run方法
        print ('staring', self.name, 'at:', ctime())
        self.res=self.func(self.args)   #python2中的用法，apply(self.func, self.args)
        print( self.name, 'finished at:', ctime())
        

class MainWindow(QMainWindow, Ui_MainWindow, threading.Thread):
    """
    Class documentation goes here.
    """
    def __init__(self, parent=None):
        """
        Constructor
        自动导入两个类：Ui_MainWindow、MainWindow
        两个方法：初始化方法、setupUi方法
        @param parent reference to the parent widget
        @type QWidget
        """
        super(MainWindow, self).__init__(parent)
        threading.Thread.__init__(self)  #对类进行初始化
        
        self.setupUi(self)
        self.graphicsView.mousePressEvent=self.my_clicked
        #ASR的几个参数
        self.framerate=16000        #采样率，有时候可以设置为8000
        self.NUM_SAMPLES=2000  #采样点
        self.channels=1                #一个声道，百度的要求一个声道
        self.sampwidth=2            #两个字节十六位
        self.TIME=2                    #条件变量，可以设置定义录音的时间
        
        self.audio_file='ASR_out.wav'  #指定录音文件，ASR_out.wav
        
        self.wav_queue=Queue.Queue(1024)  #初始化wav的队列
        self.file_name_index =1  #指定文件索引

    def run(self):  #重写run方法
        print ('staring', self.name, 'at:', ctime())  
        self.res=self.my_record()    #让录音先一直运行self.func换为my_record
        print( self.name, 'finished at:', ctime())
        
    def writeQ(self, queue, data):
        queue.put(data, 1)
    def readQ(self, queue):
        val=queue.get(1)
        return val

    #ASR的五个函数
    def save_wave_file(self,filename,data):
        '''save the date to the wavfile'''
        wf=wave.open(filename,'wb')  #二进制写入模式
        wf.setnchannels(self.channels)     
        wf.setsampwidth(self.sampwidth)   #两个字节16位
        wf.setframerate(self.framerate)       #帧速率
        wf.writeframes(b"".join(data))  #把数据加进去，就会存到硬盘上去wf.writeframes(b"".join(data)) 
        wf.close()
    def my_record(self):
        while True:
            pa=PyAudio()
            stream=pa.open(format = paInt16,channels=1,
                           rate=self.framerate,input=True,
                           frames_per_buffer=self.NUM_SAMPLES)
            my_buf=[]
            count=0
            print('.')
            while count<self.TIME* 10:#控制录音时间
                string_audio_data = stream.read(self.NUM_SAMPLES) #每读完2000个采样加1
                my_buf.append(string_audio_data)
                count+=1
                print('当前正在录音(同时录制系统内部和麦克风的声音)……')
            print('录制结束！')
            
            if my_buf:  #如果有数据就继续处理，因为输出含有b''，my_buf !=b''
                if self.file_name_index <11:  #10个文件循环线程进行(录音等)
                    pass
                else:
                    self.file_name_index =1
                filename=str(self.file_name_index)+'.wav'
                self.save_wave_file(filename=filename, data=my_buf)
                self.writeQ(queue=self.wav_queue, data=filename) #把文件名写到一个队列里面，初始化程序的时候列也要初始化
                self.file_name_index +=1
                print(filename, 'saved')
            else:
                print('file not saved')
                
    #        self.save_wave_file('ASR_out.wav',my_buf)
            stream.close()
    def play(self):
        chunk=2014
        wf=wave.open(self.audio_file,'rb')  #audio_file='16k.wav'  #指定录音文件,wave.open(r"16k.wav",'rb')
        p=PyAudio()
        stream=p.open(format=p.get_format_from_width(wf.getsampwidth()),channels=
        wf.getnchannels(),rate=wf.getframerate(),output=True)
        while True:
            data=wf.readframes(chunk)
            if data ==b"":break
            stream.write(data)
        stream.stop_stream()   # 停止数据流
        stream.close()
        p.terminate() 
        print('play函数结束！')
        
    def get_token(self):
        server = "https://openapi.baidu.com/oauth/2.0/token?"
        grant_type = "client_credentials"
        #API Key
        client_id = "qRHV7hrxj8vtAGuZOpG0zW58"
        #Secret Key
        client_secret = "Bg3Bmx3uPeCUnvuHxnS16HLnNiVPuPnz" 

        #拼url
        url ="%sgrant_type=%s&client_id=%s&client_secret=%s"%(server,grant_type,client_id,client_secret)
        #获取token
        res = requests.post(url)
        token = json.loads(res.text)["access_token"]
        print('get_token调用函数成功！'+token)
        return token
    def audio2txt(self, token):   #get_word(token)
        while True:     #不停的去读wav文件
            if self.wav_queue.qsize():  #如不等于0就
                filename=self.readQ(queue=self.wav_queue)
            else:
                continue
            with open(self.audio_file, "rb") as f:
                speech = base64.b64encode(f.read()).decode('utf8')
            size = os.path.getsize(self.audio_file)
            headers = { 'Content-Type' : 'application/json'} 
            url = "https://vop.baidu.com/server_api"
            data={
                    "format":"wav",
                    "rate":"16000",
                    "dev_pid":"1536",
                    "speech":speech,
                    "cuid":"niu_play",
                    "len":size,
                    "channel":1,
                    "token":token,
                }
            req = requests.post(url,json.dumps(data),headers)
            result = json.loads(req.text)
            print(result)
            ret0=result["result"][0]
            print(ret0)
            result1="The content of the identified audio file is"+str(ret0)
            self.textBrowser.append(result1) #向多行文本添加内容
            spk = win32com.client.Dispatch("SAPI.SpVoice")  
            spk.Speak(ret0)
            sleep(0.3)
            return result 

        

        
    def my_clicked(self, e):
        print('自定义的点击事件函数')
        webbrowser.open('www.baidu.com') #默认浏览器打开指定网址

    @pyqtSlot()
    def on_pushButton_clicked(self):
        """
        Slot documentation goes here.
        """
        print(self.textBrowser.toPlainText()) #获取多行文本框得内容
    
    @pyqtSlot()
    def on_pushButton_2_clicked(self):
        """
        Slot documentation goes here.
        这是测试按钮的槽函数
        """
        self.lineEdit.setText( "")  #清除单行文本框内容
        self.lineEdit_2.setText( "")
        self.textBrowser.setText( "") #清除多行文本框内容
        print("清除登陆账号信息！")
    
    @pyqtSlot()
    def on_pushButton_3_clicked(self):
        """
        Slot documentation goes here.
        """
        my_str=self.lineEdit.text()+":"+self.lineEdit_2.text() #获取单行文本框的内容
        self.textBrowser.append(my_str) #向多行文本添加内容
        print(my_str)
        Button3=QMessageBox.information(self, u'提示信息框', u'输入的账号和密码已经存储到数据库中！', )
        
    @pyqtSlot()
    def on_pushButton_4_clicked(self):
        """
        Slot documentation goes here.
        该按钮命令调用的预定义的函数，退出的命令
        """

    @pyqtSlot()
    def on_pushButton_5_clicked(self):
        """
        Slot documentation goes here.
        """
        Button5=QMessageBox.question(self, u'提示信息框', u'是否全部保存到数据库中？')
        if Button5==0:
            print('全部保存中……')
        else:
            print('没有保存')
            
        
    @pyqtSlot()
    def on_pushButton_6_clicked(self):
        """
        Slot documentation goes here.
        """
        Button6=QMessageBox.warning(self, u'提警告信息框', u'没有警告信息，请继续输入！')
        
    @pyqtSlot()
    def on_pushButton_7_clicked(self):
        """
        Slot documentation goes here.
        """
        Button7=QMessageBox.critical(self, u'严重警告！', u'没有严重警告信息，请继续输入！')
        
    @pyqtSlot()
    def on_pushButton_8_clicked(self):
        """
        Slot documentation goes here.
        """
        Button8=QMessageBox.about(self, u'AI简介', u'AI更多知识，可以百度《一个处女座的程序猿》查看相关CSDN博客……')

    @pyqtSlot()
    def on_pushButton_9_clicked(self):
        """
        Slot documentation goes here.
        """
        self.graphicsView.setStyleSheet("border-image: url(:/im/image/AI (4).jpg);")
        
        

 
#    @pyqtSlot(QUrl)
    def on_textBrowser_anchorClicked(self, p0):
        """
        Slot documentation goes here.
        
        @param p0 DESCRIPTION
        @type QUrl
        """

    @pyqtSlot()
    def on_radioButton_clicked(self):
        """
        Slot documentation goes here.
        """
        print('同时选择其他三个首个radioButton')
        self.radioButton_12.setChecked(True)
        self.radioButton_16.setChecked(True)
        self.radioButton_20.setChecked(True)
        self.label_4.setStyleSheet("border-image: url(:/im/image/AI (4).jpg);")
    
    @pyqtSlot()
    def on_radioButton_2_clicked(self):
        """
        Slot documentation goes here.
        """
    
    @pyqtSlot()
    def on_radioButton_4_clicked(self):
        """
        Slot documentation goes here.
        """
        
    
    @pyqtSlot()
    def on_radioButton_3_clicked(self):
        """
        Slot documentation goes here.
        """

    
    @pyqtSlot()
    def on_radioButton_12_clicked(self):
        """
        Slot documentation goes here.
        """

    
    @pyqtSlot()
    def on_radioButton_13_clicked(self):
        """
        Slot documentation goes here.
        """

    
    @pyqtSlot()
    def on_radioButton_14_clicked(self):
        """
        Slot documentation goes here.
        """

    
    @pyqtSlot()
    def on_radioButton_15_clicked(self):
        """
        Slot documentation goes here.
        """

    
    @pyqtSlot()
    def on_radioButton_16_clicked(self):
        """
        Slot documentation goes here.
        """

    
    @pyqtSlot()
    def on_radioButton_17_clicked(self):
        """
        Slot documentation goes here.
        """
        self.label_4.setStyleSheet("border-image: url(:/im/image/Bigdata.jpg);")
        

    
    @pyqtSlot()
    def on_radioButton_18_clicked(self):
        """
        Slot documentation goes here.
        """

    
    @pyqtSlot()
    def on_radioButton_19_clicked(self):
        """
        Slot documentation goes here.
        """
        self.label_4.setStyleSheet("border-image: url(:/im/image/Iot.jpg);")

    
    @pyqtSlot()
    def on_radioButton_20_clicked(self):
        """
        Slot documentation goes here.
        """

    
    @pyqtSlot()
    def on_radioButton_21_clicked(self):
        """
        Slot documentation goes here.
        """

    
    @pyqtSlot()
    def on_radioButton_22_clicked(self):
        """
        Slot documentation goes here.
        """

    
    @pyqtSlot()
    def on_radioButton_23_clicked(self):
        """
        Slot documentation goes here.
        """
    
    @pyqtSlot(int)
    def on_dial_valueChanged(self, value):
        """
        Slot documentation goes here.
        
        @param value DESCRIPTION
        @type int
        """
        self.lcdNumber.display(value)
        print(value)

    @pyqtSlot(int)
    def on_horizontalSlider_valueChanged(self, value):
        """
        Slot documentation goes here.
        
        @param value DESCRIPTION
        @type int
        """
        self.lcdNumber.display(value)
    
    @pyqtSlot(int)
    def on_verticalSlider_valueChanged(self, value):
        """
        Slot documentation goes here.
        
        @param value DESCRIPTION
        @type int
        """
        self.lcdNumber.display(value)
    
    @pyqtSlot()
    def on_pushButton_10_clicked(self):
        """
        Slot documentation goes here.
        """
        #会返回两个参数(输入的信息，窗体执行状态信息)
        my_str, ok=QInputDialog.getText(self, u'String', u'输入提示', QLineEdit.Normal, u'请在此处输入你的名字')
        print(my_str)
        self.textBrowser.append('NAME'+':'+my_str) #向多行文本添加内容
    
    @pyqtSlot()
    def on_pushButton_11_clicked(self):
        """
        Slot documentation goes here.
        """
        #QInputDialog.getInteger    QInputDialog.getInteger
        my_str, ok=QInputDialog.getInt(self, u'Int', u'请在此处输入你的年龄', 30, 0 ,100) #默认显示30
        print(my_str)
        self.textBrowser.append('AGE'+':'+str(my_str)) #向多行文本添加内容

    
    @pyqtSlot()
    def on_pushButton_12_clicked(self):
        """
        Slot documentation goes here.
        """
        my_str, ok=QInputDialog.getDouble(self, u'Int', u'请在此处输入你的年薪(以万为单位)', 3.14, 0 ,100) #默认显示30
        print(my_str)
        self.textBrowser.append('SALARY'+':'+str(my_str)) #向多行文本添加内容
        
        
    @pyqtSlot()
    def on_pushButton_13_clicked(self):
        """
        Slot documentation goes here.
        """
        my_list=['Facebook', 'Google']
        my_list.append(u'阿里巴巴')
        my_list.append(u'腾讯')
        my_list.append(u'百度')
        my_str, ok=QInputDialog.getItem(self, u'Dropdown box',u'请选择你想进入的公司', my_list )
        self.textBrowser.append('COMPANY'+':'+my_str) #向多行文本添加内容

    @pyqtSlot()
    def on_pushButton_14_clicked(self):
        """
        Slot documentation goes here.
        """
        my_info=dialog()
        my_info.exec_()
        
    @pyqtSlot()
    def on_actionClose_3_triggered(self):
        """
        Slot documentation goes here.
        """
        sys.exit(0) #退出
        
    
    @pyqtSlot()
    def on_actionAbout_triggered(self):
        """
        Slot documentation goes here.
        """
        webbrowser.open('blog.csdn.net/qq_41185868') #默认浏览器打开指定网址
        spk.Speak(u"欢迎访问我的CSDN博客，谢谢！") 
        actionAbout=QMessageBox.about(self, u'AI简介', u'AI更多知识，可以百度《一个处女座的程序猿》查看相关CSDN博客……')

    @pyqtSlot()
    def on_actionInstruction_triggered(self):
        """
        Slot documentation goes here.
        """
        actionAbout=QMessageBox.information(self, u'Instructions', u'Artificial intelligence course software is a new online internet teaching system designed by Jason Niu.')
 

    @pyqtSlot()
    def on_actionOpen_triggered(self):
        """
        Slot documentation goes here.
        """
        my_file_path, type=QFileDialog.getOpenFileName(self, u'选择文件', '')
        print(my_file_path, type)
        
        #通过文件后缀名来判断，利用不同的方式打开doc或者txt文件
        if my_file_path[-4:]=='.doc' or my_file_path[-4:]=='.wps' or my_file_path[-5:]=='.docx':  #my_file_path[-4:]获取最后四位值，一般是获取后缀
#            my_worddoc=word.Documents.Open(my_file_path,)  #my_file_path.replace(u'1',u'\\')将文件路径中的/改为\(第一个是转义符)
#            my_count=my_worddoc.Paragraphs.Count
#            for i in range(my_count): #循环读取
#                my_pr=my_worddoc.Paragraphs[i].Range   #对于每一行传入变量i索引寻找
#                print(my_pr.text)
#            self.textBrowser.append(my_pr.text)
#            print('将doc或wps的内容成功写入到文本框中！')
#            my_worddoc.Close() #文本关闭
#            import docx
#            doc=docx.Document(my_file_path)
#            for my_paragraph in doc.paragraphs:
#                print(my_paragraph.text)
#                self.textBrowser.append(my_paragraph.text)
            from win32com import  client as wc
            word=wc.Dispatch('Kwps.Application') #启动独立的进程，调用windows系统自带的程序，比如Excel.Application等
            word.Visible=0    #后台打开，前端不必显示
            self.textBrowser.append(niu_read_docx(my_file_path))
            

        elif my_file_path[-4:]=='.txt':
            f=open(my_file_path)
            my_data=f.read()
            f.close()
            self.textBrowser.append(my_data)  #my_data.decode('gbk')
            print('将txt的内容成功写入到文本框中！')
        elif my_file_path.endswith('.xlsx'):  #i
            print('成功打开表格文件！')
#            from win32com import  client as wc
#            excel=wc.Dispatch('Ket.Application') #启动独立的进程，调用windows系统自带的程序，比如Excel.Application等
#            excel.Visible=0    #后台打开，前端不必显示
#            my_excel=excel.Workbooks.Open(my_file_path)
#            print(my_excel.Sheets.Count) #默认每个表格文件含有3个sheet
#            my_sheet=my_excel.Sheets('Sheet1') #选择表Sheet1
#            print(my_sheet.UsedRange.Rows.Count) #输出已占用内容的行数
#            print(my_sheet.UsedRange.Columns.Count)
#            #嵌套for循环输出表格内容
#            for i in range(my_sheet.UsedRange.Rows.Count):
#                for j in range(my_sheet.UsedRange.Columns.Count):
#                    if my_sheet.Cells(i+1,j+1).Value:    #if条件判断空格为1才用输出，(0即None则不输出)
#                        print(my_sheet.Cells(i+1,j+1).Value)
#            my_excel.Close()  #关闭文件
#            excel.Quit()          #关闭程序
            from xlrd import open_workbook
            wb=open_workbook(my_file_path)
            for s in wb.sheets():
                for row in range(s.nrows):
                    for col in range(s.ncols):
                        if s.cell(row, col).value:    #if条件判断空格为1才用输出，(0即空格则不输出)
                            print(s.cell(row, col).value)
                    

        else:
            QMessageBox.information(self, u'information', '不支持的文件格式' )
        
    
    @pyqtSlot()
    def on_actionOpen_2_triggered(self):
        """
        Slot documentation goes here.
        """
        my_file_path, type=QFileDialog.getOpenFileName(self, u'选择文件', '')
        print(my_file_path)
        
    @pyqtSlot()
    def on_actionSave_triggered(self):
        """
        Slot documentation goes here.
        """
        my_file_path, type=QFileDialog.getOpenFileName(self, u'保存文件', '')
        f=open(my_file_path, 'a+')  #a+是写入模式
        my_data=self.textBrowser.toPlainText()   #my_data2.decode('gbk')
        f.write(my_data)
        f.close()
        print('将文本框的内容成功写入到txt文档中！')

    
    @pyqtSlot()
    def on_actionClose_triggered(self):
        """
        Slot documentation goes here.
        """
        my_data=self.textBrowser.toPlainText()
        my_file, type=QFileDialog.getSaveFileName(self, u'文件另存为', '')
        print('文件另存为路径：'+my_file)
    
    @pyqtSlot()
    def on_actionClose_2_triggered(self):
        """
        Slot documentation goes here.
        """

    
    @pyqtSlot()
    def on_actionCopy_triggered(self):
        """
        Slot documentation goes here.
        """
    
    @pyqtSlot()
    def on_actionPaste_triggered(self):
        """
        Slot documentation goes here.
        """

    
    @pyqtSlot()
    def on_actionCut_3_triggered(self):
        """
        Slot documentation goes here.
        """

    
    @pyqtSlot()
    def on_actionUndo_triggered(self):
        """
        Slot documentation goes here.
        """

    
    @pyqtSlot()
    def on_actionRedo_triggered(self):
        """
        Slot documentation goes here.
        """

    
    @pyqtSlot()
    def on_actionDelete_triggered(self):
        """
        Slot documentation goes here.
        """

    
    @pyqtSlot()
    def on_actionSelect_All_triggered(self):
        """
        Slot documentation goes here.
        """

    
    @pyqtSlot()
    def on_actionRegister_triggered(self):
        """
        Slot documentation goes here.
        """

    
    @pyqtSlot()
    def on_actionOfficial_triggered(self):
        """
        Slot documentation goes here.
        """

    
    @pyqtSlot()
    def on_actionContact_Us_triggered(self):
        """
        Slot documentation goes here.
        """
        
        
    @pyqtSlot()
    def on_pushButton_15_clicked(self):
        """
        Slot documentation goes here.
        """

    
    @pyqtSlot()
    def on_pushButton_16_clicked(self):
        """
        Slot documentation goes here.
        """

    
    @pyqtSlot()
    def on_pushButton_17_clicked(self):
        """
        Slot documentation goes here.
        """
        self.my_record()         #进行录音并保存

    
    @pyqtSlot()
    def on_pushButton_18_clicked(self):
        """
        Slot documentation goes here.
        """
        self.play()                #播放录音文件
        
        
    @pyqtSlot()
    def on_pushButton_19_clicked(self):
        """
        Slot documentation goes here.
        """

    
    @pyqtSlot()
    def on_pushButton_20_clicked(self):
        """
        Slot documentation goes here.
        """
        self.token=self.get_token()
        ret = self.audio2txt(self.token)


        
        
#if __name__ == "__main__":
#    import sys
#    app = QtWidgets.QApplication(sys.argv)
#    ui = MainWindow()
#    ui.show()
#    sys.exit(app.exec_())
if __name__ == "__main__":  
    import sys  
    app = QtWidgets.QApplication(sys.argv)  
      
    #加载启动画面  
    splash=QSplashScreen(QPixmap(":/im/image/AI (6).jpg")) #QSplashScreen启动界面  
    splash.show()  
    
    splash.showMessage(u'正在加载初始界面……',Qt.AlignCenter,Qt.yellow) #设置文字在正中央  
    time.sleep(1)  
    splash.showMessage(u'正在加载音频资源……',Qt.AlignCenter,Qt.red)  


    spk = win32com.client.Dispatch("SAPI.SpVoice")  
    spk.Speak(u"Welcome to use Artificial intelligence designed by Jason Niu") 

    time.sleep(1)  
    splash.showMessage(u'正在渲染界面……',Qt.AlignCenter,Qt.blue)  
    time.sleep(1)  
      
    app.processEvents()  #防止加载的图片卡死启动  
  
    ui = MainWindow()  
    splash.finish(ui) #在主窗口下，启动画面结束  
    
    ui.setDaemon(True)  #在线程结束的时候子线程也会结束
    ui.start()     #继承线程thread中的start方法，开启线程
    
    #识别程序在上边的主类中实现，下边实现录音的程序record_t
    record_t=MyThread(ui.audio2txt, (ui.get_token()), ui.audio2txt.__name__)  #开启线程进行识别文本内容(传递函数地址，)
    record_t.setDaemon(True)#进程随之结束
    record_t.start()  #进程随之结束
    
    ui.show()  
    sys.exit(app.exec_())
    

